from __future__ import annotations

from pathlib import Path
from runpy import run_path
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
SYSTEM_PROMPT = run_path(str(ROOT / "prompts.py"))["SYSTEM_PROMPT"]
SOURCE = ROOT / "docs" / "バクラク事業部技術_実験報告書_完成版.docx"
OUTPUT = ROOT / "docs" / "バクラク事業部技術_実験報告書_完成版_改訂.docx"
ASSETS = ROOT / "docs" / "report_assets"
FONT = "MS Mincho"
BLACK = RGBColor(0, 0, 0)


def font_run(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("eastAsia", "ascii", "hAnsi"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph(paragraph, text: str, *, size: float | None = None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    font_run(run, size=size)


def replace_in_runs(paragraph, old: str, new: str) -> None:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            font_run(run)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def remove_table(table) -> None:
    element = table._element
    element.getparent().remove(element)


def set_cell(cell, text: str, *, bold: bool = False, size: float = 9.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    font_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, value: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def move_before(anchor, element) -> None:
    anchor._element.addprevious(element)


def add_prompt_block(doc: Document, anchor, title: str, lines: list[str], *, page_break: bool) -> None:
    heading = doc.add_paragraph(style="Heading 3")
    heading.paragraph_format.page_break_before = page_break
    heading.paragraph_format.keep_with_next = True
    set_paragraph(heading, title, size=11.0)
    move_before(anchor, heading._element)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    row = table.rows[0]
    add_no_split(row)
    cell = row.cells[0]
    cell.width = Inches(6.25)
    set_cell_margins(cell, 80)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(8.5)
    paragraph.paragraph_format.keep_together = True
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        font_run(run, size=7.5)
        if index != len(lines) - 1:
            run.add_break()
    move_before(anchor, table._element)


def image_relationship_id(paragraph) -> str:
    blips = paragraph._element.xpath(".//a:blip")
    if not blips:
        raise RuntimeError("Expected an embedded image before the caption")
    return blips[0].get(qn("r:embed"))


def replace_image_before_caption(
    doc: Document,
    caption_text: str,
    image_path: Path,
    alt_text: str,
) -> None:
    paragraphs = doc.paragraphs
    index = next(i for i, p in enumerate(paragraphs) if p.text.startswith(caption_text))
    picture = paragraphs[index - 1]
    rid = image_relationship_id(picture)
    doc.part.rels[rid].target_part._blob = image_path.read_bytes()
    for doc_property in picture._element.xpath(".//wp:docPr"):
        doc_property.set("descr", alt_text)
        doc_property.set("title", alt_text)


def find_paragraph(doc: Document, prefix: str):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def main() -> None:
    doc = Document(SOURCE)

    # Use Japanese terminology consistently, while defining the English term once.
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph, "Active RAG", "アクティブRAG")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph, "Active RAG", "アクティブRAG")
    first_rag = next(p for p in doc.paragraphs if "必要箇所を反復検索するアクティブRAG" in p.text)
    set_paragraph(
        first_rag,
        first_rag.text.replace("アクティブRAGも", "アクティブRAG（Active RAG）も"),
    )

    # Replace embedded charts with the regenerated Gemini-based figures.
    replace_image_before_caption(
        doc,
        "図2",
        ASSETS / "figure_02_hypothesis_evolution.png",
        "全文入力、文字認識、アクティブRAGを検討し、決定論的な完全ページ選択を採用した。",
    )
    replace_image_before_caption(
        doc,
        "図4",
        ASSETS / "figure_04_benchmark_efficiency.png",
        "Gemini 3.7 Flash・中程度の同一75資料で三戦略の処理速度を比較する。",
    )
    replace_image_before_caption(
        doc,
        "図5",
        ASSETS / "figure_05_benchmark_quality.png",
        "Gemini 3.7 Flash・中程度の同一75資料で三戦略の全呼出し入力トークン数を比較する。",
    )

    set_paragraph(
        find_paragraph(doc, "更新済み全コーパスのGLM-5.3思考モードでは"),
        "更新済み全コーパスのGLM-5.3思考モードでは、必要ページだけを選ぶ戦略三が102/102資料を完了し、マクロ平均正確性99.23%、マイクロ正確性99.14%であった。Geminiで三戦略すべてが完了した同一75資料では、戦略三の全呼出し入力トークン数は戦略一比89.1%少なかった。",
    )
    set_paragraph(
        find_paragraph(doc, "全コーパス表の「平均初回入力トークン数」"),
        "全コーパス表の「平均初回入力トークン数」は、ベンチマークJSONが保持する初回抽出呼出しだけの値である。契約修復と根拠再試行は含まない。P50処理時間は、完了実行の処理時間の中央値である。一方、同一75資料の三戦略比較では、対応するprediction.jsonを照合し、初回抽出、契約修復、根拠再試行の報告入力トークンを合算した「平均全呼出し入力トークン数」を用いた。この75資料では、実行された18件の根拠再試行を含め、使用量記録の欠落はなかった。",
    )
    set_paragraph(
        find_paragraph(doc, "入力戦略だけを公平に比べるため"),
        "入力戦略だけを公平に比べるため、Gemini 3.7 Flash・中程度で三戦略すべてが完了した同一75資料を抽出した。ここではモデル、推論設定、資料をそろえている。",
    )

    strategy_tables = [t for t in doc.tables if t.rows and t.rows[0].cells[0].text == "戦略"]
    table4 = strategy_tables[3]
    values = [
        ["戦略一", "96.59%", "96.62%（1,429/1,479）", "37.87秒", "88,936"],
        ["戦略二", "96.41%", "96.42%（1,426/1,479）", "34.74秒", "89,414"],
        ["戦略三", "97.73%", "97.90%（1,448/1,479）", "29.07秒", "9,715"],
    ]
    for row, row_values in zip(table4.rows[1:], values):
        add_no_split(row)
        for cell, value in zip(row.cells, row_values):
            set_cell(cell, value)

    set_paragraph(
        find_paragraph(doc, "図4　三戦略の処理速度"),
        "図4　三戦略の処理速度（戦略一比）\nGemini 3.7 Flash・中程度で三戦略すべてが完了した同一75資料のP50処理時間から算出した。モデル間ではなく、入力戦略だけを比較する。",
        size=9.0,
    )
    set_paragraph(
        find_paragraph(doc, "図5　三戦略の全呼出し入力トークン数"),
        "図5　三戦略の全呼出し入力トークン数\n図4と同じ75資料における1資料当たりの平均である。初回抽出、契約修復、根拠再試行の報告入力トークンを合算した。戦略三は戦略一に対して89.1%削減した。",
        size=9.0,
    )
    set_paragraph(
        find_paragraph(doc, "図4の速度指数は"),
        "図4の速度指数は、戦略一のP50処理時間を1.00倍とした。戦略二は1.09倍、戦略三は1.30倍である。このGemini条件では戦略三が最速であった。図5では、戦略三の平均全呼出し入力トークン数が戦略一比89.1%少ない。採用判断は、速度だけではなく、マイクロ正確性、入力削減、読取不能資料への対応、根拠追跡を合わせて行った。",
    )
    set_paragraph(
        find_paragraph(doc, "三戦略すべてが完了した同一75資料では"),
        "Geminiで三戦略すべてが完了した同一75資料では、戦略三が最高のマイクロ正確性97.90%、最短のP50処理時間29.07秒、最少の平均全呼出し入力トークン数9,715を示した。",
    )

    # Replace the internal benchmark-manifest table with reader-facing reproduction guidance.
    heading_65 = find_paragraph(doc, "6.5 再現条件")
    set_paragraph(heading_65, "6.5 再現・試用方法", size=14.0)
    caption_65 = find_paragraph(doc, "表6　ベンチマーク再現条件")
    remove_paragraph(caption_65)
    repro_table = next(t for t in doc.tables if t.rows[0].cells[0].text == "ベンチマーク再現条件")
    remove_table(repro_table)
    anchor_66 = find_paragraph(doc, "6.6 評価上の注意")
    repro_paragraphs = [
        "本実装は、READMEの「Getting started」に従ってローカル環境で再現できる。Python 3.11以上とNode.js 20以上を用意し、PythonとReactの依存関係を導入する。.env.exampleを.envへ複製し、LLM_PROVIDER、LLM_API_KEY、LLM_MODELなどを設定する。APIキーは環境変数としてサーバー側だけで読み込み、ブラウザ、ログ、実行結果、ソースコードへ書き出さない。Flask APIと構築済み画面はhttp://localhost:5000で起動する。React画面を開発する場合は、Flaskとは別にVite開発サーバーを起動する。正確なコマンドと検証手順はREADMEを参照する。",
        "公開試用環境はhttps://assignment.mohamedfuad.comである。試用用のOpenRouter設定はAWS上のバックエンドに保持し、APIキーをブラウザへ配布しない。残クレジットの範囲では、目安として10〜20件のPDFを投入し、三戦略の出力を比較できる。利用可能件数は、PDFの長さ、選択した戦略、モデル、および残クレジットによって変動する。",
        "ベンチマークでは、正解表を推論入力に含めない。PDFと共通プロンプトだけで抽出を完了した後、独立した評価工程で正解表を参照する。GPT-5.6 Lunaは、現行プロジェクトに呼出し可能なプロバイダ経路とモデルIDがないため、本版の比較表へ含めていない。同一102資料・三戦略で追試するには、利用可能なAPIエンドポイントと正式なモデルIDを先に確定する必要がある。",
    ]
    for text in repro_paragraphs:
        paragraph = anchor_66.insert_paragraph_before()
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.space_after = Pt(6)
        set_paragraph(paragraph, text)

    # Replace the abbreviated prompt with the complete runtime SYSTEM_PROMPT.
    appendix_intro = find_paragraph(doc, "以下は、実験条件を再現するために必要な指示文の主要部分")
    set_paragraph(
        appendix_intro,
        "以下に、prompts.pyから実行時に生成されるSYSTEM_PROMPT全文を、そのまま掲載する。省略、要約、正解値の挿入は行っていない。論理区分ごとにページを分け、表の途中で分断しない。",
    )
    short_prompt = next(t for t in doc.tables if t.rows[0].cells[0].text.startswith("Use only the Annual Report"))
    remove_table(short_prompt)
    source_intro = find_paragraph(doc, "戦略三の中心となる処理は")
    lines = SYSTEM_PROMPT.strip().splitlines()
    ranges = [(0, 26), (26, 51), (51, 81), (81, 103), (103, 124), (124, len(lines))]
    titles = [
        "D.1 SYSTEM_PROMPT全文（前提・規則 1/2）",
        "D.2 SYSTEM_PROMPT全文（規則 2/2）",
        "D.3 SYSTEM_PROMPT全文（項目対応 1/3）",
        "D.4 SYSTEM_PROMPT全文（項目対応 2/3）",
        "D.5 SYSTEM_PROMPT全文（項目対応 3/3）",
        "D.6 SYSTEM_PROMPT全文（出力契約）",
    ]
    for index, ((start, end), title) in enumerate(zip(ranges, titles)):
        add_prompt_block(doc, source_intro, title, lines[start:end], page_break=index > 0)
    code_heading = doc.add_paragraph(style="Heading 3")
    code_heading.paragraph_format.page_break_before = True
    code_heading.paragraph_format.keep_with_next = True
    set_paragraph(code_heading, "D.7 主要ソースコード", size=11.0)
    move_before(source_intro, code_heading._element)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
