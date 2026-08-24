from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPORT = Path(__file__).with_name("バクラク事業部技術_実験報告書_完成版.docx")


def _set_page_break_before(paragraph, enabled: bool) -> None:
    paragraph.paragraph_format.page_break_before = enabled


def _section_break(template, *, landscape: bool):
    section = deepcopy(template)
    section_type = section.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        page_size_index = next(
            (index for index, child in enumerate(section) if child.tag == qn("w:pgSz")),
            len(section),
        )
        section.insert(page_size_index, section_type)
    section_type.set(qn("w:val"), "nextPage")

    page_size = section.find(qn("w:pgSz"))
    if landscape:
        page_size.set(qn("w:w"), "16834")
        page_size.set(qn("w:h"), "11909")
        page_size.set(qn("w:orient"), "landscape")
        margins = section.find(qn("w:pgMar"))
        for side, value in {
            "top": "620",
            "right": "500",
            "bottom": "620",
            "left": "500",
            "header": "360",
            "footer": "360",
        }.items():
            margins.set(qn(f"w:{side}"), value)
    return section


def _attach_section_break(paragraph, section) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()
    existing = paragraph_properties.find(qn("w:sectPr"))
    if existing is not None:
        paragraph_properties.remove(existing)
    paragraph_properties.append(section)


def _set_table_widths(table, widths) -> None:
    table.autofit = False
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(width.twips))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:w"), str(width.twips))
            cell_width.set(qn("w:type"), "dxa")


def _format_prompt_as_landscape_page(document: Document) -> None:
    heading = next(p for p in document.paragraphs if p.text.startswith("D.1 SYSTEM_PROMPT"))
    code_heading = next(p for p in document.paragraphs if p.text.startswith("D.2 主要ソースコード"))
    prompt_table = next(
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text.startswith("You are a financial-document extraction system.")
    )
    body_section = document._element.body.sectPr

    # End the preceding portrait section immediately before D.1.
    intro_element = heading._element.getprevious()
    if intro_element is None or intro_element.tag != qn("w:p"):
        raise RuntimeError("Prompt introduction paragraph was not found")
    intro_paragraph = next(p for p in document.paragraphs if p._element is intro_element)
    _attach_section_break(intro_paragraph, _section_break(body_section, landscape=False))
    _set_page_break_before(heading, False)

    # End the dedicated landscape section after the prompt table. Reuse the
    # existing break paragraph when the finalizer is run more than once.
    next_element = prompt_table._element.getnext()
    if next_element is not None and next_element.tag == qn("w:p") and next_element is not code_heading._element:
        break_paragraph = next(p for p in document.paragraphs if p._element is next_element)
    else:
        break_element = OxmlElement("w:p")
        prompt_table._element.addnext(break_element)
        break_paragraph = next(p for p in document.paragraphs if p._element is break_element)
    break_paragraph.paragraph_format.space_before = Pt(0)
    break_paragraph.paragraph_format.space_after = Pt(0)
    break_paragraph.paragraph_format.line_spacing = Pt(1)
    _attach_section_break(break_paragraph, _section_break(body_section, landscape=True))
    _set_page_break_before(code_heading, False)

    # The full prompt remains exact, but the landscape width supports a
    # materially larger type size than the former 5-point portrait layout.
    _set_table_widths(prompt_table, [Inches(3.42), Inches(3.42), Inches(3.42)])
    for cell in prompt_table.rows[0].cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = cell_properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            cell_properties.append(shading)
        shading.set(qn("w:fill"), "F7F7F7")
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(7.5)
            paragraph.paragraph_format.keep_together = True
            for run in paragraph.runs:
                run.font.size = Pt(6.5)


def main() -> None:
    document = Document(REPORT)
    toc_pages = {
        "要旨": "3",
        "1.0 背景・課題設定": "4",
        "2.0 目的・仮説・評価観点": "5",
        "3.0 対象範囲・前提条件": "6",
        "4.0 実施方法／三戦略": "7",
        "5.0 実装・最終アーキテクチャ": "8",
        "6.0 評価方法・結果": "10",
        "7.0 考察": "14",
        "8.0 結論・次の課題": "16",
        "参考文献": "17",
        "付録": "18",
    }
    for paragraph in document.paragraphs:
        if not paragraph.style.name.lower().startswith("toc"):
            continue
        heading = paragraph.text.split("\t", 1)[0]
        if heading not in toc_pages:
            continue
        text_nodes = paragraph._element.xpath(".//w:t")
        text_nodes[-1].text = toc_pages[heading]
    _format_prompt_as_landscape_page(document)
    appendix_e = next(p for p in document.paragraphs if p.text.startswith("付録E"))
    appendix_e.paragraph_format.page_break_before = True
    trailing = document.paragraphs[-1]
    following = trailing._element.getnext()
    if not trailing.text and following is not None and following.tag == qn("w:sectPr"):
        trailing._element.getparent().remove(trailing._element)
    document.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
