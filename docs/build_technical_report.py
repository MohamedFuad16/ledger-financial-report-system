from __future__ import annotations

import re
from runpy import run_path
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path("/Users/mfuad16/Downloads/バクラク事業部技術 実験報告書.docx")
REPORT_SOURCE = ROOT / "docs" / "technical_experiment_report_final_ja.md"
OUT = ROOT / "docs" / "バクラク事業部技術_実験報告書_完成版.docx"
ASSETS = ROOT / "docs" / "report_assets"

JP_FONT = "MS Mincho"
BODY_COLOR = RGBColor(0, 0, 0)
PALE_BLUE = "F2F2F2"
PALE_GRAY = "F5F5F5"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None, italic: bool | None = None) -> None:
    run.font.name = JP_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), JP_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), JP_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), JP_FONT)
    lang = run._element.get_or_add_rPr().find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.get_or_add_rPr().append(lang)
    lang.set(qn("w:val"), "ja-JP")
    lang.set(qn("w:eastAsia"), "ja-JP")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=100, bottom=60, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    specs = {
        "Normal": (10.5, False, BODY_COLOR),
        "Normal (Web)": (10.5, False, BODY_COLOR),
        "Title": (28, True, BODY_COLOR),
        "Heading 1": (18, True, BODY_COLOR),
        "Heading 2": (14, True, BODY_COLOR),
        "Heading 3": (12, True, BODY_COLOR),
        "Caption": (9, False, BODY_COLOR),
        "List Bullet": (10.5, False, BODY_COLOR),
        "List Number": (10.5, False, BODY_COLOR),
        "Hyperlink": (10.5, False, BODY_COLOR),
        "TOC 1": (10.5, False, BODY_COLOR),
        "TOC 2": (9.5, False, BODY_COLOR),
        "TOC Heading": (20, True, BODY_COLOR),
    }
    for name, (size, bold, color) in specs.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = JP_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), JP_FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), JP_FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), JP_FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        if not hasattr(style, "paragraph_format"):
            continue
        pf = style.paragraph_format
        if name.startswith("Heading"):
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
        elif name in {"Normal", "Normal (Web)"}:
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pf.space_after = Pt(5)
        elif name.startswith("TOC "):
            pf.space_before = Pt(0)
            pf.space_after = Pt(3)


def insert_native_toc_after(paragraph) -> None:
    toc_p = OxmlElement("w:p")
    toc_p_pr = OxmlElement("w:pPr")
    toc_p.append(toc_p_pr)
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    run.append(begin)
    toc_p.append(run)
    run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-1" \\h \\z \\u '
    run.append(instruction)
    toc_p.append(run)
    run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run.append(separate)
    placeholder = OxmlElement("w:t")
    placeholder.text = "目次を更新してください。"
    run.append(placeholder)
    toc_p.append(run)
    run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run.append(end)
    toc_p.append(run)
    paragraph._p.addnext(toc_p)


def normalize_existing_front_matter(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            size = run.font.size.pt if run.font.size else None
            set_run_font(run, size=size, color=BODY_COLOR)
        # python-docx omits runs nested inside hyperlinks from paragraph.runs.
        # Normalize those runs directly so the cover remains black throughout.
        for run_el in paragraph._p.xpath(".//w:hyperlink/w:r"):
            r_pr = run_el.get_or_add_rPr()
            fonts = r_pr.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                r_pr.insert(0, fonts)
            for attr in ("ascii", "eastAsia", "hAnsi"):
                fonts.set(qn(f"w:{attr}"), JP_FONT)
            color = r_pr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                r_pr.append(color)
            color.set(qn("w:val"), "000000")
            color.attrib.pop(qn("w:themeColor"), None)

    replacements = {
        "２０２６年８月２１日": "２０２６年８月２４日",
        "課題概要": "要旨",
        "本課題では、３Ｍ社の２０２２年年次報告書を一つの検証事例として使用し、年次報告書から貸借対照表の資産項目および金額を抽出する仕組みを構築します。":
            "本報告書の目的は、3M社の2022年年次報告書から、指定された27項目の資産表を作成する仕組みを設計し、他社資料への適用可能性を評価することである。",
        "特定の企業や報告書の形式だけに依存するのではなく、他企業の年次報告書が入力された場合にも対応できる、汎用的な仕組みを目標とします。":
            "課題を、入力表現、指示と出力形式、検証と評価の三つへ分解した。文字認識を使わない全文入力、文字認識後の全文入力、必要ページだけを選ぶ入力の三戦略を比較した。",
        "本報告書では、課題に対する初期の考え方、各実験で設定した仮説、実装方法、試行錯誤の過程、実験結果、精度評価、発見した技術的課題、および各課題に対する改善の見通しを記録します。":
            "更新済み全コーパスのGLM-5.3思考モードでは、必要ページだけを選ぶ戦略三が102/102資料を完了し、マクロ平均正確性99.23%、マイクロ正確性99.14%であった。三戦略すべてが完了した同一75資料では、戦略三の全呼出し入力トークン数は戦略一比90.1%少なかった。",
        "現在の進捗": "主要結果",
        "複数の仮説に基づく手法を実装し、比較実験・精度検証を実施中":
            "最終的に、決定論的なページ選択を採用した。",
        "現在、年次報告書から必要な財務情報を正確に特定・抽出し、指定された貸借対照表の形式へ変換するため、複数のアプローチを比較しています。":
            "採用理由は、精度を監視しながら入力を大幅に減らし、選択した根拠を再現できるためである。ただし、すべての条件で最高精度または最短時間になるわけではない。",
        "（実験結果の確定に合わせ、来週より本報告書を順次執筆・更新予定です。）":
            "結論として、事業PDFから信頼できる構造化データを作るには、意味写像を大規模言語モデルへ任せ、入力選択、形式、算術、評価を決定論的な処理で管理する構成が適する。",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text in replacements:
            replacement = replacements[paragraph.text]
            paragraph.clear()
            run = paragraph.add_run(replacement)
            set_run_font(run, size=10.5, color=BODY_COLOR)
            if replacement == "要旨":
                # Keep the linked TOC as a dedicated second page.
                paragraph.paragraph_format.page_break_before = True
            elif replacement == "主要結果":
                # The abstract title is Heading 1; its first subsection must be
                # Heading 2 so the semantic outline does not skip a level.
                paragraph.style = doc.styles["Heading 2"]

    # The supplied cover uses spacer paragraphs sized for its original font.
    # Removing the final empty spacers prevents the cover from spilling onto an
    # otherwise blank second page after Japanese font normalization.
    toc = next((p for p in doc.paragraphs if p.text == "目次"), None)
    if toc is not None:
        # Keep the TOC title clear of the printable area's top edge in both
        # Word and PDF exports.  An explicit paragraph gap is more reliable
        # here than depending on the template's first-line position.
        toc.paragraph_format.page_break_before = True
        toc.paragraph_format.space_before = Pt(18)
        toc.paragraph_format.space_after = Pt(12)
        previous = toc._p.getprevious()
        while previous is not None and previous.tag == qn("w:p"):
            if "".join(previous.itertext()).strip():
                break
            candidate = previous
            previous = previous.getprevious()
            candidate.getparent().remove(candidate)

    # Replace the original one-row index with a native, linked Word TOC field.
    if doc.tables:
        old_toc_table = doc.tables[0]._tbl
        old_toc_table.getparent().remove(old_toc_table)
    if toc is not None:
        insert_native_toc_after(toc)
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_portrait_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    return section


def add_inline_figure(doc: Document, image_path: Path, title: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    figure_width = 6.70 if image_path.name == "figure_03_final_architecture.png" else 6.45
    shape = p.add_run().add_picture(str(image_path), width=Inches(figure_width))
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(1)
    cap.paragraph_format.space_after = Pt(5)
    run = cap.add_run(title)
    set_run_font(run, size=9, bold=True, color=BODY_COLOR)
    run = cap.add_run("\n" + caption)
    set_run_font(run, size=8.5, color=BODY_COLOR)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.page_break_before = (
            text.startswith(("1.0 ", "2.0 ", "6.0 "))
            or text == "付録"
        )
        p.paragraph_format.space_before = Pt(0)
        pPr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)
        pPr.append(borders)
    for run in p.runs:
        set_run_font(run)


def add_body_paragraph(doc: Document, text: str, *, quote: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    if quote:
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.right_indent = Inches(0.18)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), PALE_BLUE)
        pPr.append(shd)
    # Render Markdown bold spans without carrying Markdown punctuation.
    pos = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > pos:
            r = p.add_run(text[pos:match.start()])
            set_run_font(r, size=10.5, color=BODY_COLOR)
        r = p.add_run(match.group(1))
        set_run_font(r, size=10.5, bold=True, color=BODY_COLOR)
        pos = match.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run_font(r, size=10.5, color=BODY_COLOR)


def add_list_item(doc: Document, text: str, numbered: bool = False) -> None:
    # Use authored numbers rather than Word's continuing list state. This keeps
    # separate lists (for example, boundaries and next actions) at 1-based order.
    p = doc.add_paragraph(style="Normal" if numbered else "List Bullet")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.5)
    r = p.add_run(text)
    set_run_font(r, size=10.3, color=BODY_COLOR)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F3F4")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("\n".join(lines))
    set_run_font(r, size=8.3, color=BODY_COLOR)
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    row_pr.append(cant_split)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    title = rows[0][0] if rows[0] else ""
    table_id_match = re.match(r"^(表(?:\d+[a-z]?|[A-Z]\d+))\s*", title)
    if table_id_match:
        table_id = table_id_match.group(1)
        cap = doc.add_paragraph(style="Caption")
        cap.paragraph_format.keep_with_next = True
        cap.paragraph_format.space_before = Pt(5)
        cap.paragraph_format.space_after = Pt(3)
        run = cap.add_run(title)
        set_run_font(run, size=9, bold=True, color=BODY_COLOR)
        rows = [row[:] for row in rows]
        first_header_overrides = {
            "表3a": "戦略",
            "表3b": "戦略",
            "表3c": "戦略",
            "表3d": "戦略",
            "表3e": "区分",
            "表2a": "比較観点",
            "表4": "戦略",
        }
        rows[0][0] = first_header_overrides.get(
            table_id,
            re.sub(r"^表(?:\d+[a-z]?|[A-Z]\d+)\s*", "", title),
        )
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", text))
            if i == 0:
                set_cell_shading(cell, "E6E6E6")
                set_run_font(r, size=8.5, bold=True, color=BODY_COLOR)
            else:
                if i % 2 == 0:
                    set_cell_shading(cell, PALE_GRAY)
                set_run_font(r, size=8.2, color=BODY_COLOR)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    tr_pr.append(repeat_header)
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        row_pr.append(cant_split)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells)


def append_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    # Existing Word draft supplies the cover, TOC, and abstract.
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1.0"))
    lines = lines[start:]
    i = 0
    code_mode = False
    code_lines: list[str] = []
    inserted_problem = False
    inserted_hypothesis = False
    inserted_architecture = False
    pending_figure: tuple[Path, str, str] | None = None
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if code_mode:
                add_code_block(doc, code_lines)
                code_lines = []
                code_mode = False
            else:
                code_mode = True
            i += 1
            continue
        if code_mode:
            code_lines.append(line)
            i += 1
            continue
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped == "[[FIGURE_04_BENCHMARK_EFFICIENCY]]":
            add_inline_figure(
                doc,
                ASSETS / "figure_04_benchmark_efficiency.png",
                "図4　三戦略の処理速度（戦略一比）",
                "Gemini 3.7 Flash・中程度で三戦略すべてが完了した同一75資料のP50処理時間から算出した。モデル間ではなく、入力戦略だけを比較する。",
            )
            i += 1
            continue
        if stripped == "[[FIGURE_05_BENCHMARK_QUALITY]]":
            add_inline_figure(
                doc,
                ASSETS / "figure_05_benchmark_quality.png",
                "図5　三戦略の全呼出し入力トークン数",
                "図4と同じ75資料における1資料当たりの平均である。初回抽出、契約修復、根拠再試行の報告入力トークンを合算した。戦略三は戦略一に対して89.1%削減した。",
            )
            i += 1
            continue
        if stripped == "[[FULL_SYSTEM_PROMPT]]":
            system_prompt = run_path(str(ROOT / "prompts.py"))["SYSTEM_PROMPT"]
            add_code_block(doc, system_prompt.strip().splitlines())
            i += 1
            continue
        if stripped.startswith("|"):
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not is_separator_row(cells):
                    table_rows.append(cells)
                i += 1
            add_markdown_table(doc, table_rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            hashes, text = heading.groups()
            level = min(3, len(hashes) - 1)
            add_heading(doc, text, level)
            if text.startswith("1.0") and not inserted_problem:
                pending_figure = (
                    ASSETS / "figure_01_problem_decomposition.png",
                    "図1　課題を三つの境界へ分解する",
                    "課題を、入力表現、指示と出力形式、検証と評価の三つへ分解する。",
                )
                inserted_problem = True
            if text.startswith("2.2") and not inserted_hypothesis:
                pending_figure = (
                    ASSETS / "figure_02_hypothesis_evolution.png",
                    "図2　仮説の変化と採用判断",
                    "全文入力、文字認識、アクティブRAGを検討し、決定論的な完全ページ選択を採用した。",
                )
                inserted_hypothesis = True
            if text.startswith("5.2") and not inserted_architecture:
                pending_figure = (
                    ASSETS / "figure_03_final_architecture.png",
                    "図3　最終アーキテクチャ",
                    "PDF解析基盤で文字層を確認し、必要ページだけを文字認識してから、根拠ページを意味写像と検証へ渡す。",
                )
                inserted_architecture = True
            i += 1
            continue
        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            add_list_item(doc, f"{m.group(1)}. {m.group(2)}", numbered=True)
            i += 1
            continue
        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:])
            i += 1
            continue
        if stripped.startswith("> "):
            add_body_paragraph(doc, stripped[2:], quote=True)
            i += 1
            continue
        # Join consecutive prose lines into one paragraph.
        prose = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", "|", "- ", "> ", "```"))
                    or re.match(r"^\d+\.\s+", nxt)):
                break
            prose.append(nxt)
            i += 1
        add_body_paragraph(doc, " ".join(prose))
        if pending_figure is not None:
            add_inline_figure(doc, *pending_figure)
            pending_figure = None


def add_running_header(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
        if not p.text.strip():
            p.text = "バクラク事業部技術　実験報告書"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            set_run_font(run, size=8, color=BODY_COLOR)


def main() -> None:
    doc = Document(SOURCE)
    configure_styles(doc)
    normalize_existing_front_matter(doc)
    # Separate the supplied three-page front matter from the completed report body.
    add_portrait_section(doc)
    append_markdown(doc, REPORT_SOURCE.read_text(encoding="utf-8"))
    add_running_header(doc)

    props = doc.core_properties
    props.title = "バクラク事業部技術 実験報告書"
    props.subject = "Annual Reportから資産の部27項目を抽出するシステムの設計・実験・コード監査"
    props.author = "モハメド・フアド"
    props.comments = "2026年8月24日版。製品コードは読み取り専用監査のみ。"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
